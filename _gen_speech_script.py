# -*- coding: utf-8 -*-
"""
Generate the English speaker script (Word .docx) for the MSDM6980 defense slides.
Slide structure follows MSDM6980_Beamer.tex (29 frames, 7 sections).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:/Users/marcozhu/Desktop/6980/MSDM6980_Defense_Script.docx"

# (slide_no, section, title, est_time_sec, script_paragraphs[])
SLIDES = [
    # ---------- Title & TOC ----------
    (1, "Opening", "Title Page", 45, [
        "Good afternoon, Professors. My name is ZHU TIANYI, student number 21229582, and I am very honoured to present my MSDM6980 capstone today. My supervisor is Mr. Hanson WONG.",
        "The title of my project is \"Machine Learning for Non-Performing Asset Repayment Probability Prediction and Strategy Optimization.\" It is a comparative study between traditional machine-learning models and frontier large language models on a real, highly imbalanced Hong Kong unsecured-debt portfolio.",
        "Over the next twenty minutes I will walk you through the problem, my methodology, the empirical results, the resulting collection strategy, and a head-to-head benchmark of twenty-six LLM configurations against specialised ML.",
    ]),
    (2, "Opening", "Presentation Outline", 30, [
        "The talk is organised in seven parts. I will first motivate the problem and state my research objectives. Then I will summarise my own contributions, before describing the data and the modelling pipeline.",
        "Part four reports the ML results; part five turns those probabilities into an optimised collection strategy. Part six is the LLM benchmark, which is the most novel piece of this project. I will close with conclusions, limitations and future work.",
    ]),

    # ---------- I. Introduction ----------
    (3, "I. Introduction", "Background and Problem Statement", 75, [
        "Non-performing assets are a major and growing risk for global financial institutions. The portfolio I work with contains sixteen thousand and forty-eight delinquent unsecured-debt accounts from Hong Kong.",
        "The operational challenge is simple to state but hard to solve: collection capacity is finite. We have agent calls, auto-dialer calls, and SMS or email channels, each with very different unit costs. We must decide, account by account, which channel to spend on whom.",
        "Statistically, this is a binary prediction problem under extreme class imbalance. Only nine point one five percent of accounts ever repay within three years. That means a naive model that predicts \"no one repays\" is already ninety-one percent accurate but completely useless.",
        "The strategic transition we want to enable is from uniform treatment, where every debtor gets the same script, to precision targeting, where calibrated probabilities drive the channel choice.",
    ]),
    (4, "I. Introduction", "Research Objectives (O1 -- O4)", 70, [
        "I framed the project around four concrete objectives.",
        "Objective one: develop and compare four machine-learning models — Logistic Regression, Random Forest, XGBoost, and a custom MLP version two — for binary repayment prediction.",
        "Objective two: evaluate them on three independent dimensions — discrimination through AUC, calibration through the Brier score, and economic outcome through net recovery and return on investment.",
        "Objective three: turn those probabilities into a four-tier collection strategy under explicit cost and capacity constraints.",
        "Objective four, and the most novel contribution: systematically benchmark twenty-six frontier LLM configurations — zero-shot, few-shot, and fine-tuned — against the specialised ML baselines.",
    ]),

    # ---------- II. Contributions ----------
    (5, "II. Contributions", "Contribution 1: End-to-End ML Pipeline", 70, [
        "My first contribution is a fully reproducible end-to-end ML pipeline tailored to this domain.",
        "The most important methodological choice is the missing-value treatment. The variable \"months since last payment\" has one thousand one hundred and sixty-five missing values. These are not random — they correspond to debtors who have never paid the original creditor at all. Mean or median imputation would erase a strong predictive signal, so I encode them with a sentinel value plus a binary \"never paid\" flag.",
        "I tailored the encoding to each model family — one-hot plus z-score for the linear and neural models, ordinal encoding for the tree models. For probability calibration, I reserve a dedicated twenty-percent calibration fold and run five-fold out-of-fold Platt scaling.",
        "Finally, to make the work usable by non-technical stakeholders, I built a nine-tab interactive HTML dashboard.",
    ]),
    (6, "II. Contributions", "Contribution 2: Strategy & Frontier LLM Benchmark", 80, [
        "My second contribution sits in the strategy and LLM-benchmark layer.",
        "I evaluate twenty-six configurations spanning OpenAI o1 and o3, Claude 4 Opus, Gemini 2.5 Pro, DeepSeek-R1, Qwen3, Llama 4 and Grok 3, plus three fine-tuned variants — across zero-shot, few-shot of ten examples, and full fine-tuning.",
        "A key finding here is that prompt design is itself a hyperparameter. My domain-calibrated chain-of-thought prompt injects the nine-point-one-five-percent base rate, the explicit cost grid, and known monotonicity priors. This single change lifted Claude 3.7 Sonnet from AUC zero-point-six-two-five to zero-point-six-seven-one, and cut the model's self-reported ROI error from over two hundred percent down to under fifteen percent.",
        "On the strategy side, I built a four-tier allocation module with explicit cost and capacity constraints, and along the way I uncovered a striking pattern: an inverse balance-to-repayment relationship — sub-two-hundred-HKD accounts repay around twenty-five percent, while accounts above two hundred thousand HKD repay essentially zero percent.",
    ]),

    # ---------- III. Methodology ----------
    (7, "III. Data and Methodology", "Dataset Schema", 70, [
        "Let me briefly walk you through the dataset. We have sixteen thousand and forty-eight accounts described by nine features.",
        "Categorical features include loan type, district, and a multiple-account flag. Ordinal features include the balance bucket. Numerical features cover borrower birth year, account vintage, time since last activity, time since last payment, and the outstanding balance proxy.",
        "The target is whether the account repays within three years — a positive rate of nine point one five percent.",
        "I split the data seventy-five-twenty-five into a training set M of twelve thousand thirty-six accounts and a hold-out test set T of four thousand and twelve accounts. All evaluation numbers I will show are on T, the untouched test set.",
    ]),
    (8, "III. Data and Methodology", "Preprocessing Pipeline", 60, [
        "The preprocessing pipeline has four steps.",
        "First, sentinel flagging — I replace nulls in the last-payment-date variable with a plus-nine-nine-nine sentinel and add a binary \"never paid\" flag.",
        "Second, encoding — one-hot for the linear and neural models, ordinal for the tree models.",
        "Third, scaling — I z-score features only for the MLP. Tree models are scale-invariant, so I leave them on the raw scale.",
        "Fourth, splitting — inside the training set M, I create a sixty-twenty-twenty split for training, validation, and a calibration fold. The calibration fold is reserved exclusively for Platt scaling, never seen during training.",
    ]),
    (9, "III. Data and Methodology", "Traditional Models: LR, RF, XGBoost", 65, [
        "I implement three traditional models as baselines.",
        "Logistic Regression is the linear baseline, with L2 regularisation, regularisation strength C equals ten, and balanced class weights.",
        "Random Forest uses five hundred trees, square-root-d feature subsampling per split, a maximum depth of twelve, and a minimum of eight samples per leaf.",
        "XGBoost is the champion ranker. It optimises the standard regularised loss shown on the slide, with a positive-class weight of about nine point one eight to compensate for imbalance. I use two hundred trees, a learning rate of zero point zero five, max depth six, and column subsampling at zero point eight five.",
    ]),
    (10, "III. Data and Methodology", "Advanced Architecture: MLP v2", 75, [
        "The deep-learning model — MLP version two — was deliberately upgraded after version one underperformed. Three design changes mattered.",
        "First, I introduced a residual block: each hidden state is added back to its non-linear transform under a Swish activation. This stabilises gradient flow on a shallow tabular network.",
        "Second, I added an explicit feature-interaction layer of sixty-four dimensions, which materialises pairwise products. Tabular models often fail because they cannot represent feature interactions cheaply, and this layer addresses exactly that.",
        "The full architecture is: linear two-fifty-six, Swish, residual block of one-twenty-eight, the feature-interaction layer, linear thirty-two, dropout zero point three, and the output head. Training uses AdamW with a OneCycle learning-rate schedule, label smoothing of zero point zero three, gradient clipping at three, for two hundred epochs.",
        "After this upgrade, the MLP reaches AUC zero point seven zero six three — competitive with XGBoost on calibration, and clearly better than the version-one baseline.",
    ]),
    (11, "III. Data and Methodology", "Probability Calibration: Platt Scaling", 65, [
        "Calibration deserves its own slide because it is what makes the strategy economically valid.",
        "XGBoost outputs margin scores, not probabilities. The raw scores can rank accounts well but they systematically over- or under-estimate the absolute repayment probability. That breaks any threshold-based allocation rule.",
        "I correct this with Platt scaling, fitting a logistic transformation on the dedicated five-fold out-of-fold calibration set. The two parameters A and B are estimated by maximum likelihood.",
        "This step is essential because the action tier — Agent, Dialer, SMS, or write-off — depends on the absolute probability, not just the rank order.",
    ]),

    # ---------- IV. Results ----------
    (12, "IV. Results and Analysis", "ML Model Performance (Test Set)", 90, [
        "Now the core empirical result. This is the test-set performance table on four thousand and twelve held-out accounts.",
        "On AUC, XGBoost wins at zero point seven three zero five — that is the best ranking.",
        "On Brier score, Logistic Regression wins at zero point zero eight four two — that is the best calibration.",
        "On net recovery and ROI, Logistic Regression wins again at one point four six eight million HKD and twenty-nine point four times ROI.",
        "The headline insight is therefore counter-intuitive but important: XGBoost wins ranking, but Logistic Regression wins economics. Why? Because the economic objective is not \"who do I rank highest\" but \"what is the absolute probability that this debtor pays back\". Logistic Regression's outputs are already well-calibrated probabilities, while XGBoost needs Platt scaling to recover that property.",
        "This has a direct production implication: in a deployed system, ranking power and calibration must both be measured. A model that wins one and loses the other can still lose money.",
    ]),
    (13, "IV. Results and Analysis", "Discrimination vs. Calibration vs. Economics", 60, [
        "This three-panel figure visualises that result. Panel (a) is ROC-AUC, where XGBoost leads. Panel (b) is Brier score, where Logistic Regression leads — and lower is better. Panel (c) is net recovery, where Logistic Regression also leads.",
        "What the panels collectively show is that no single metric is sufficient. The choice of model depends on what we optimise for. For our profit objective, calibration matters as much as discrimination.",
    ]),
    (14, "IV. Results and Analysis", "ROC Curve Analysis", 45, [
        "The ROC curves confirm the AUC ranking visually. XGBoost reaches AUC zero point seven three zero five, with all four models clearly above the zero-point-five random-baseline diagonal. Random Forest and the MLP are close behind, and Logistic Regression sits last on AUC but, as we just saw, first on profit.",
    ]),
    (15, "IV. Results and Analysis", "Operational Accuracy (Confusion Matrices)", 55, [
        "The confusion matrices give the operational view. Among the four models, XGBoost yields the fewest false positives — meaning fewer accounts are misclassified as likely repayers and therefore fewer expensive Agent-tier calls are wasted on debtors who will never pay. This is a direct cost saving in production.",
    ]),
    (16, "IV. Results and Analysis", "Permutation Feature Importance", 65, [
        "The permutation feature-importance analysis on XGBoost reveals what is actually driving the predictions.",
        "Borrower birth year dominates by a wide margin, with a delta-AUC of zero point one three zero. Older cohorts have distinct repayment dynamics — a fact that is consistent with both regulatory experience and our own monotonicity prior.",
        "District contributes about zero point zero four one, and the balance bucket about zero point zero three three. These three features alone explain the bulk of the model's discriminative power.",
    ]),
    (17, "IV. Results and Analysis", "Inverse Balance--Repayment Relationship", 75, [
        "This is one of the most actionable findings. There is a strong, almost monotonic, inverse relationship between outstanding balance and repayment probability.",
        "Sub-two-hundred-HKD accounts repay at roughly twenty-five percent — about three times the portfolio average. These are typically debtors whose balances are below the psychological threshold of \"a meaningful debt\", and they often clear them when contacted.",
        "On the other end, accounts with balances above two hundred thousand HKD repay at roughly zero percent. For this tier, collection effort is value-destroying — every dollar spent contacting them is wasted.",
        "The strategic implication is to concentrate cheap channels — Dialer and SMS — on the high-volume, low-balance population, and to write off the high-balance tail.",
    ]),

    # ---------- V. Strategy ----------
    (18, "V. Strategy", "Optimized Queue Allocation Summary", 85, [
        "These probabilities now feed an optimised four-tier collection strategy.",
        "Eighteen percent of accounts — seven hundred and twenty-two — go to the Agent tier. They are the highest-probability segment and justify the most expensive channel at eighty-five HKD per contact, returning seventeen point seven times ROI.",
        "Forty-two percent — about one thousand six hundred and eighty-five accounts — go to Auto-Dialer at twelve HKD per contact. ROI here climbs to forty-nine times because the cost per contact drops while the underlying repayment probability is still meaningful.",
        "Thirty percent go to SMS or email at one point five HKD per contact. The ROI peaks at one hundred and five point six times — this is the cheapest, most efficient tier.",
        "The remaining ten percent — four hundred and two accounts — are written off because the expected recovery does not even cover the cheapest contact cost.",
        "Bottom line: a contact budget of eighty-three thousand HKD generates an expected net recovery of two point two seven million HKD, an overall portfolio ROI of twenty-seven point two times.",
    ]),
    (19, "V. Strategy", "Strategy Concentration (Pareto Effect)", 50, [
        "This figure shows the Pareto effect explicitly. The top twenty percent of accounts — ranked by predicted probability — capture fifty point six percent of the expected net recovery.",
        "That is the entire economic justification for risk-based segmentation: half of the value sits in twenty percent of the book, and ML lets us find that twenty percent.",
    ]),

    # ---------- VI. LLM Benchmark ----------
    (20, "VI. Frontier LLM Benchmark", "Benchmarking AI-as-Judge (26 LLM Configurations)", 85, [
        "Now I turn to the LLM benchmark. The motivation is a simple practical question: given how strong frontier reasoning models have become, can we just prompt a model like Claude 4 Opus instead of training XGBoost?",
        "I evaluate twenty-six configurations across three regimes and four model families.",
        "Heuristics include random guessing and a hand-crafted rule-based scorer.",
        "Pre-reasoning zero-shot includes GPT-4o, Claude 3.5 Sonnet, Gemini 1.5 Pro, DeepSeek-V3, and Qwen2.5.",
        "Reasoning zero-shot-plus-thinking includes OpenAI o1 and o3, Claude 3.7 and 4 Sonnet, Claude 4 Opus, Gemini 2.5 Pro, DeepSeek-R1, Qwen3, Llama 4, and Grok 3.",
        "Few-shot of ten and fine-tuned variants include GPT-4o FS, o3 FS, Claude 4 Sonnet FS, plus three fine-tuned models — Llama-3.3-70B, Qwen3-32B, and GPT-4.1-mini.",
        "On the protocol: eight configurations were run live against the API on a five-hundred-account stratified sub-sample. The remaining eighteen rows are public-benchmark-adjusted using TabLLM, CARTE and FinBench, with a plus-or-minus zero-point-zero-one-five AUC band that I report transparently in the paper.",
    ]),
    (21, "VI. Frontier LLM Benchmark", "Domain-Calibrated CoT Prompt", 75, [
        "Before showing the numbers, I want to spend a slide on the prompt itself, because prompt design dominated the variance I saw across runs.",
        "The role frame casts the model as a financial risk analyst evaluating delinquent unsecured-debt accounts in Hong Kong.",
        "The priors block injects three pieces of domain knowledge: the portfolio base rate of nine point one five percent, the explicit cost grid for the three channels, and two monotonicity priors — older borrowers repay more, and balance is inversely related to repayment.",
        "The reasoning steps force the model to first reason from features, then pick a tier, then estimate a net recovery and ROI consistent with the cost grid. The output schema is strict JSON.",
        "The effect was significant: Claude 3.7 Sonnet's AUC moved from zero point six two five to zero point six seven one, and the self-reported ROI error collapsed from over two hundred percent to under fifteen percent. So calibrated prompts are not optional — they are part of the modelling work.",
    ]),
    (22, "VI. Frontier LLM Benchmark", "Headline: Best LLM per Family vs. ML Anchors", 90, [
        "Here is the headline LLM-versus-ML table. I show only the best representative from each regime, plus the ML anchors. The full twenty-six-row version is in Appendix E of the paper.",
        "Random guessing gives AUC zero point five and zero net recovery, as expected.",
        "The rule-based heuristic reaches zero point five six two — useful as a sanity floor.",
        "The best pre-reasoning zero-shot model — GPT-4o — reaches zero point six two eight.",
        "The best reasoning zero-shot — Claude 4 Opus — climbs to zero point six nine four. That is genuinely impressive for an out-of-the-box generalist model.",
        "Few-shot with ten examples on Claude 4 Sonnet adds about half a point — to zero point seven zero one.",
        "Fine-tuning is where things get interesting. Fine-tuned GPT-4.1-mini reaches AUC zero point seven two two, beating Random Forest and matching Logistic Regression on AUC.",
        "But on the right-hand side of the table — the economic columns — the picture changes. Logistic Regression still delivers one point four six eight million HKD net recovery and twenty-nine times ROI. The best fine-tuned LLM only reaches one point one four four million — twenty-eight percent below LR. The reason, again, is calibration.",
    ]),
    (23, "VI. Frontier LLM Benchmark", "LLM-vs-ML Comparison (May 2026)", 60, [
        "This four-panel figure visualises the same story. Red bars are LLMs, blue bars are ML.",
        "Panel (a) shows the AUC ranking: the LLMs catch up but specialised XGBoost still leads.",
        "Panel (b) shows the Brier score, where LR is clearly the calibration winner.",
        "Panel (c) shows net recovery — and you can see the LLMs cluster well below the ML anchors.",
        "Panel (d) is the AUC-versus-ROI Pareto frontier: ML occupies the upper-right region — high discrimination AND high economic return — while LLMs sit on a clearly inferior frontier.",
    ]),
    (24, "VI. Frontier LLM Benchmark", "Five Findings (F1--F5)", 100, [
        "I distil the LLM benchmark into five findings.",
        "Finding one: the new floor is the old ceiling. Pre-reasoning zero-shot models now sit at AUC zero point six zero to zero point six four. Reasoning models with thinking sit at zero point six six to zero point seven zero. The level of capability that required heavy prompting two years ago is now the default.",
        "Finding two: few-shot is now redundant for reasoning models. Going from o3 zero-shot at zero point six eight one to o3 few-shot of ten at zero point seven zero one only adds two AUC points. This is a meaningful change in how we should think about LLM cost-of-deployment.",
        "Finding three: fine-tuning closes the AUC gap. Fine-tuned GPT-4.1-mini reaches zero point seven two two — Random Forest parity, Logistic Regression parity.",
        "Finding four: ML still wins decisively on net recovery. Logistic Regression delivers one point four six eight million HKD versus one point one four four million for the best LLM — a twenty-eight percent gap. The bottleneck is calibration, not discrimination.",
        "Finding five: operational economics still favour ML by a factor of about seventy thousand. LR scores an account in zero point zero five milliseconds. Claude 4 Opus takes about three thousand five hundred milliseconds. Scoring the full sixteen-thousand-account portfolio takes under one second on ML versus around fifteen hours and roughly eight US dollars on the strongest reasoning LLM.",
    ]),
    (25, "VI. Frontier LLM Benchmark", "Operational Analysis: The AI Tax", 60, [
        "This figure makes the operational case visible. I call it the \"AI tax\".",
        "On latency, ML is at sub-millisecond scale. Frontier LLMs are at two to three thousand milliseconds — four orders of magnitude slower.",
        "On cost, API calls make LLM inference roughly one thousand times more expensive at portfolio scale.",
        "On explainability, ML retains a strict advantage for regulatory and audit reporting, because feature attributions are stable, deterministic, and reproducible.",
    ]),
    (26, "VI. Frontier LLM Benchmark", "Deployment Recommendation", 75, [
        "Given those findings, my deployment recommendation is a hybrid architecture.",
        "ML is the production scorer for all sixteen thousand and forty-eight accounts — specifically a calibrated Logistic Regression and XGBoost ensemble. This handles the high-throughput, high-stakes scoring loop.",
        "Frontier reasoning LLMs — Claude 4 Opus and o3 — are reserved for the borderline tier, the roughly eight percent of accounts whose calibrated probability sits between zero point one eight and zero point three two. These are precisely the cases where a small change in probability flips the channel decision and changes the economic outcome.",
        "Two reasons make this trade attractive on the borderline set: the natural-language rationales support compliance and audit review, and the marginal cost is acceptable because we are only invoking the LLM on a small subset.",
    ]),

    # ---------- VII. Conclusion ----------
    (27, "VII. Conclusion", "Summary of Principal Findings", 80, [
        "To summarise the principal findings.",
        "First, ranking and profit are not the same objective. XGBoost is the best ranker at AUC zero point seven three zero five. Logistic Regression is the best profit-maker, with one point four six eight million HKD net recovery and a twenty-nine times ROI, thanks to its superior native calibration.",
        "Second, calibration is mandatory in any threshold-based allocation. Without Platt scaling, our four-tier strategy literally does not work.",
        "Third, the dominant drivers are age and balance. Resources should shift toward low-balance, high-volume tiers, and away from the high-balance tail.",
        "Fourth, frontier LLMs narrow but do not close the gap. Fine-tuning achieves AUC parity, but trails Logistic Regression by twenty-eight percent on net recovery, and reasoning LLMs are uneconomical at portfolio scale.",
        "And fifth: specialised ML remains the gold standard for high-throughput tabular financial prediction.",
    ]),
    (28, "VII. Conclusion", "Limitations and Future Work", 65, [
        "Five honest limitations and matching future-work directions.",
        "One: temporal validation. The current split is random; a vintage-cohort time-series cross-validation is needed before deployment.",
        "Two: survival analysis. Predict not just whether a debtor will pay, but when — that opens up time-discounted cash-flow optimisation.",
        "Three: macroeconomic covariates. Integrating Hong Kong unemployment, HIBOR, and CPI signals should improve robustness across business cycles.",
        "Four: a live A/B test. The tiered strategy must be field-tested against the heuristic baseline.",
        "Five: a hybrid ML-plus-LLM workflow. ML for prediction, LLM for borderline-case adjudication and customer-facing explanations.",
    ]),
    (29, "Closing", "Thank You", 25, [
        "That concludes the presentation. Thank you very much for your time and attention. I am happy to take any questions.",
    ]),
]


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Default font: Calibri 11
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Cover ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MSDM6980 Capstone Defense — Speaker Script")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x00, 0x38, 0x74)  # HKUST blue

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Machine Learning for Non-Performing Asset Repayment Probability "
                    "Prediction and Strategy Optimization")
    r.italic = True
    r.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Student: ZHU TIANYI (21229582)   |   Supervisor: Mr. Hanson WONG\n"
        "The Hong Kong University of Science and Technology   |   May 2026"
    )

    # Notes
    note = doc.add_paragraph()
    note.add_run("How to use this script. ").bold = True
    note.add_run(
        "The script is organised slide-by-slide, matching the 29 frames of "
        "MSDM6980_Beamer.tex. Each slide is annotated with its section, title, "
        "and an estimated reading time at a steady pace (~150 wpm). Total target "
        "duration: ~20 minutes plus Q&A. Read aloud naturally; the bracketed "
        "[PAUSE] markers indicate brief pauses for slide transitions."
    )

    total = sum(s[3] for s in SLIDES)
    summary = doc.add_paragraph()
    summary.add_run(f"Estimated total speaking time: ~{total // 60} min {total % 60} sec.").italic = True

    doc.add_page_break()

    # ---- Per-slide script ----
    last_section = None
    for slide_no, sect, title, secs, paragraphs in SLIDES:
        if sect != last_section:
            h1 = doc.add_paragraph()
            r1 = h1.add_run(sect)
            r1.bold = True
            r1.font.size = Pt(15)
            r1.font.color.rgb = RGBColor(0x00, 0x38, 0x74)
            last_section = sect

        h2 = doc.add_paragraph()
        r2 = h2.add_run(f"Slide {slide_no:02d}  —  {title}")
        r2.bold = True
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0xA7, 0x83, 0x37)  # HKUST gold

        meta = doc.add_paragraph()
        meta.add_run(f"Estimated time: ~{secs} sec").italic = True

        for p in paragraphs:
            para = doc.add_paragraph(p)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing = 1.3

        pause = doc.add_paragraph("[PAUSE — advance slide]")
        pause.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pause.runs[0].italic = True
        pause.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.add_paragraph()  # spacer

    doc.save(OUT)
    print(f"WROTE: {OUT}")


if __name__ == "__main__":
    main()
